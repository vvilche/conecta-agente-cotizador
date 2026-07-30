import sys
import io
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest
import docx
from src.operations.official_word_quote_builder import OfficialWordQuoteBuilder, format_currency


def test_docx_generation_and_6_headings():
    payload = {
        "partner_id": "Transelec S.A.",
        "amount_untaxed": 10000000.0,
        "amount_tax": 1900000.0,
        "amount_total": 11900000.0,
        "margin_analysis": {"boosted_margin_pct": 54.8},
        "order_line": [
            {
                "item_code": "HW-ORION-LX",
                "name": "RTU NovaTech Orion LX+",
                "product_uom_qty": 2,
                "price_unit": 3500000,
                "price_subtotal": 7000000,
            },
            {
                "item_code": "HW-BELDEN-SW",
                "name": "Switch Belden Hirschmann",
                "product_uom_qty": 1,
                "price_unit": 3000000,
                "price_subtotal": 3000000,
            },
        ],
    }

    docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(payload)
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0

    doc = docx.Document(io.BytesIO(docx_bytes))

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("1. DETALLE DE LOS SUMINISTROS Y SERVICIOS" in h for h in headings)
    assert any("2. DETALLE DE PRECIO OFERTA BASE" in h for h in headings)
    assert any("3. EXCLUSIONES DE LA OFERTA" in h for h in headings)
    assert any("4. VALIDEZ DE LA OFERTA" in h for h in headings)
    assert any("5. CONDICIONES DE PAGO" in h for h in headings)
    assert any("6. TÉRMINOS Y CONDICIONES (T&C)" in h for h in headings)


def test_metadata_block_and_defaults():
    payload = {"partner_id": "Enel Chile S.A.", "order_line": []}
    docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(payload)
    doc = docx.Document(io.BytesIO(docx_bytes))

    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "Nuestra Ref.:" in full_text
    assert "Rev 0" in full_text
    assert "Santiago, " in full_text
    assert "Enel Chile S.A." in full_text


def test_custom_reference_code_and_date():
    payload = {
        "partner_id": "Comasa S.A.",
        "reference_number": "REF-CUSTOM-2026-V1",
        "date": "Santiago, 15 de Agosto de 2026",
        "order_line": [],
    }
    docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(payload)
    doc = docx.Document(io.BytesIO(docx_bytes))

    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "REF-CUSTOM-2026-V1" in full_text
    assert "Santiago, 15 de Agosto de 2026" in full_text


def test_summary_table_6_columns_and_currency():
    for curr in ["CLP", "USD", "UF"]:
        payload = {
            "partner_id": "Engie Energía",
            "currency": curr,
            "amount_untaxed": 50000.0,
            "amount_tax": 9500.0,
            "amount_total": 59500.0,
            "order_line": [
                {
                    "item_code": "PMU-VIZIMAX",
                    "name": "PMU SynchroTeq",
                    "product_uom_qty": 1,
                    "price_unit": 50000.0,
                    "price_subtotal": 50000.0,
                }
            ],
        }
        docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(payload)
        doc = docx.Document(io.BytesIO(docx_bytes))

        table = doc.tables[0]
        assert len(table.columns) == 6

        hdr_row = table.rows[0]
        hdr_texts = [cell.text for cell in hdr_row.cells]
        assert hdr_texts[0] == "Ítem"
        assert hdr_texts[1] == "Código Partida"
        assert hdr_texts[2] == "Descripción de la Partida"
        assert hdr_texts[3] == "Cant."
        assert f"Precio Unit. Neto {curr}" in hdr_texts[4]
        assert f"Subtotal Venta {curr}" in hdr_texts[5]

        data_row = table.rows[1]
        assert data_row.cells[0].text == "1"
        assert data_row.cells[1].text == "PMU-VIZIMAX"
        assert data_row.cells[2].text == "PMU SynchroTeq"
        assert data_row.cells[3].text == "1"
        assert curr in data_row.cells[4].text
        assert curr in data_row.cells[5].text


def test_customizable_sections():
    payload = {
        "exclusions": "Custom Exclusions: Excluye montaje solar en techo.",
        "validity": "La oferta vence en 60 días corridos.",
        "payment_conditions": "50% anticipo, 50% contra entrega.",
        "terms_and_conditions": "Garantía extendida de 24 meses.",
        "order_line": [],
    }
    docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(payload)
    doc = docx.Document(io.BytesIO(docx_bytes))

    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "Custom Exclusions: Excluye montaje solar en techo." in full_text
    assert "La oferta vence en 60 días corridos." in full_text
    assert "50% anticipo, 50% contra entrega." in full_text
    assert "Garantía extendida de 24 meses." in full_text


# =====================================================================
# EXPANDED EDGE CASE & PARAMETERIZED TEST SUITE FOR WORD QUOTE BUILDER
# =====================================================================

@pytest.mark.parametrize("val, currency, expected_substring", [
    (1250.75, "UF", "1.250,75 UF"),
    (3500000, "CLP", "$ 3.500.000 CLP"),
    (4500.50, "USD", "$ 4.500,50 USD"),
    (0.0, "CLP", "$ 0 CLP"),
    (0.0, "UF", "0,00 UF"),
    (0.0, "USD", "$ 0,00 USD"),
])
def test_format_currency_helper(val, currency, expected_substring):
    formatted = format_currency(val, currency=currency)
    assert expected_substring in formatted


@pytest.mark.parametrize("currency_code", ["UF", "CLP", "USD"])
def test_word_builder_uf_and_multi_currency_totals(currency_code):
    payload = {
        "partner_id": "Colbún S.A.",
        "currency": currency_code,
        "amount_untaxed": 1000.0 if currency_code == "UF" else 1000000.0,
        "amount_tax": 190.0 if currency_code == "UF" else 190000.0,
        "amount_total": 1190.0 if currency_code == "UF" else 1190000.0,
        "order_line": [
            {
                "item_code": "ING-DIgSILENT",
                "name": "Estudio de Coordinación EDAC",
                "product_uom_qty": 1,
                "price_unit": 1000.0 if currency_code == "UF" else 1000000.0,
                "price_subtotal": 1000.0 if currency_code == "UF" else 1000000.0,
            }
        ]
    }
    docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(payload)
    doc = docx.Document(io.BytesIO(docx_bytes))
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert currency_code in full_text


@pytest.mark.parametrize("num_items", [1, 3, 5, 10])
def test_multi_item_table_row_count_and_styling(num_items):
    lines = [
        {
            "item_code": f"ITEM-{i+1:03d}",
            "name": f"Equipo o Servicio de Prueba {i+1}",
            "product_uom_qty": i + 1,
            "price_unit": 100000.0 * (i + 1),
            "price_subtotal": 100000.0 * (i + 1) * (i + 1),
        }
        for i in range(num_items)
    ]
    payload = {
        "partner_id": "Transelec S.A.",
        "order_line": lines
    }
    docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(payload)
    doc = docx.Document(io.BytesIO(docx_bytes))
    table = doc.tables[0]

    # Row count: 1 header row + num_items data rows + 3 totals rows (Subtotal, IVA, Total)
    assert len(table.rows) == 1 + num_items + 3


@pytest.mark.parametrize("missing_field", [
    "partner_id",
    "amount_untaxed",
    "amount_tax",
    "amount_total",
    "margin_analysis",
    "order_line"
])
def test_missing_payload_fields_resilience(missing_field):
    base_payload = {
        "partner_id": "Cliente Test S.A.",
        "amount_untaxed": 5000000.0,
        "amount_tax": 950000.0,
        "amount_total": 5950000.0,
        "margin_analysis": {"boosted_margin_pct": 54.8},
        "order_line": [{"item_code": "TEST", "name": "Item Test", "product_uom_qty": 1, "price_unit": 5000000, "price_subtotal": 5000000}]
    }
    del base_payload[missing_field]

    # Should build without raising exceptions
    docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(base_payload)
    assert len(docx_bytes) > 0


@pytest.mark.parametrize("custom_date", [
    "Santiago, 01 de Enero de 2026",
    "Concepción, 30 de Noviembre de 2026",
    "Antofagasta, 15 de Octubre de 2027"
])
def test_dynamic_date_injection(custom_date):
    payload = {
        "partner_id": "AES Andes",
        "date": custom_date,
        "order_line": []
    }
    docx_bytes = OfficialWordQuoteBuilder.build_quote_docx_bytes(payload)
    doc = docx.Document(io.BytesIO(docx_bytes))
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert custom_date in full_text
