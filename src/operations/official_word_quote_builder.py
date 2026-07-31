"""
Official Word (.docx) Commercial Proposal Builder for Conecta Ingeniería S.A.
Generates formatted Microsoft Word proposal documents following exact historical Conecta S.A. sections:
1. DETALLE DE LOS SUMINISTROS Y SERVICIOS
2. DETALLE DE PRECIO OFERTA BASE
3. EXCLUSIONES DE LA OFERTA
4. VALIDEZ DE LA OFERTA
5. CONDICIONES DE PAGO
6. TÉRMINOS Y CONDICIONES (T&C)
"""
import io
import datetime
from typing import Dict, Any, List
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SPANISH_MONTHS = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio",
    7: "Julio", 8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}


def set_cell_background(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def format_currency(amount: float, currency: str = "CLP") -> str:
    curr = (currency or "CLP").upper()
    val = float(amount or 0.0)
    if curr == "CLP":
        formatted_num = f"{int(round(val)):,}".replace(",", ".")
        return f"$ {formatted_num} CLP"
    elif curr in ("USD", "UF"):
        parts = f"{val:.2f}".split(".")
        integer_part = f"{int(parts[0]):,}".replace(",", ".")
        decimal_part = parts[1]
        formatted_num = f"{integer_part},{decimal_part}"
        if curr == "USD":
            return f"$ {formatted_num} USD"
        else:
            return f"{formatted_num} UF"
    else:
        return f"{val:,.2f} {curr}"


class OfficialWordQuoteBuilder:
    @staticmethod
    def build_quote_docx_bytes(payload: Dict[str, Any]) -> bytes:
        doc = docx.Document()

        # Page setup - Margins 1 inch
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        client_name = payload.get("partner_id") or payload.get("client_name") or "Cliente Coordinado Conecta"
        amount_untaxed = float(payload.get("amount_untaxed", 58628319))
        amount_tax = float(payload.get("amount_tax", round(amount_untaxed * 0.19, 0)))
        amount_total = float(payload.get("amount_total", amount_untaxed + amount_tax))
        margin_pct = float(payload.get("margin_analysis", {}).get("boosted_margin_pct", 54.8))
        lines = payload.get("order_line", [])
        currency = (payload.get("currency") or "CLP").upper()

        # Dynamic Reference Code: Default YYMMDD Rev X
        now = datetime.datetime.now()
        yymmdd = now.strftime("%y%m%d")
        rev = str(payload.get("revision", "0"))
        default_ref = f"{yymmdd} Rev {rev}"
        ref_code = (
            payload.get("reference_number")
            or payload.get("reference_code")
            or payload.get("ref_code")
            or payload.get("nuestra_ref")
            or default_ref
        )

        # Dynamic Date Formatting
        if payload.get("date") or payload.get("fecha"):
            date_str = str(payload.get("date") or payload.get("fecha"))
        else:
            month_name = SPANISH_MONTHS.get(now.month, "Julio")
            date_str = f"Santiago, {now.day} de {month_name} de {now.year}"

        attention = payload.get("attention") or "Departamento de Ingeniería & Proyectos / Gerencia de Operaciones"
        subject_ref = payload.get("subject_ref") or f"PROPUESTA COMERCIAL INTEGRAL DE AUTOMATIZACIÓN OT — {client_name}"

        # Corporate Colors
        blue_color = RGBColor(0x1E, 0x3A, 0x8A)
        muted_color = RGBColor(0x47, 0x55, 0x69)

        # Header Title
        p_title = doc.add_paragraph()
        run_title = p_title.add_run("CONECTA INGENIERÍA S.A.")
        run_title.font.name = "Calibri"
        run_title.font.size = Pt(22)
        run_title.font.bold = True
        run_title.font.color.rgb = blue_color

        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run("OFERTA TÉCNICO-COMERCIAL OFICIAL — SISTEMAS DE SUBESTACIÓN & OT")
        run_sub.font.name = "Calibri"
        run_sub.font.size = Pt(11)
        run_sub.font.bold = True
        run_sub.font.color.rgb = muted_color

        # Info Box Block
        p_meta = doc.add_paragraph()
        p_meta.add_run("Nuestra Ref.: ").bold = True
        p_meta.add_run(f"{ref_code}\n")
        p_meta.add_run("Fecha: ").bold = True
        p_meta.add_run(f"{date_str}\n")
        p_meta.add_run("Señores: ").bold = True
        p_meta.add_run(f"{client_name}\n")
        p_meta.add_run("Atención: ").bold = True
        p_meta.add_run(f"{attention}\n")
        p_meta.add_run("Ref.: ").bold = True
        p_meta.add_run(f"{subject_ref}")

        # Letter Greeting
        p_body = doc.add_paragraph()
        p_body.paragraph_format.space_before = Pt(12)
        p_body.paragraph_format.space_after = Pt(12)
        greeting_text = payload.get("greeting") or (
            "Estimados Señores:\n\n"
            "Basados en vuestros requerimientos técnicos y normativos del Sistema Eléctrico Nacional (SEN), "
            "tenemos el placer de presentar nuestra oferta para el suministro, pre-kitting en taller, "
            "configuración de comunicaciones e integración de hardware industrial en subestación."
        )
        p_body.add_run(greeting_text)

        # 1. DETALLE DE LOS SUMINISTROS Y SERVICIOS
        h1 = doc.add_heading("1. DETALLE DE LOS SUMINISTROS Y SERVICIOS", level=1)
        h1.runs[0].font.color.rgb = blue_color

        custom_supplies = payload.get("supplies_services") or payload.get("bullet_points")
        if custom_supplies:
            if isinstance(custom_supplies, list):
                for item in custom_supplies:
                    bp = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, tuple) and len(item) == 2:
                        r_bt = bp.add_run(item[0])
                        r_bt.bold = True
                        bp.add_run(item[1])
                    else:
                        bp.add_run(str(item))
            else:
                doc.add_paragraph(str(custom_supplies))
        else:
            default_bullet_points = [
                ("Equipamiento de Control & Telemetría: ", "Suministro de hardware certificado Belden Hirschmann, VIZIMAX SynchroTeq Plus PMU y NovaTech Orion LX+."),
                ("Pre-kitting Estandarizado en Taller (KittingEngine): ", "Tableros pre-cableados en taller Conecta S.A., reduciendo la estadía en terreno de 5 días a solo 1.5 días."),
                ("Banco de Pruebas HIL FAT/SAT (FatSatSimulator): ", "Simulación HIL de laboratorio para validar tramas DNP3.0 y C37.118 antes del despacho."),
                ("Tramitación Regulatoria CEN / SEC (DocAutomator): ", "Elaboración en 3 segundos de Informe IPES Puesta en Servicio y Protocolo AT-SITR-1."),
                ("Acreditación Express de Personal (AccreditationAutomator): ", "Dossier digital F30-1, ex. médicos y EPP para plataformas Sicop / Pronexo.")
            ]
            for b_title, b_desc in default_bullet_points:
                bp = doc.add_paragraph(style='List Bullet')
                r_bt = bp.add_run(b_title)
                r_bt.bold = True
                bp.add_run(b_desc)

        # 2. DETALLE DE PRECIO OFERTA BASE
        h2 = doc.add_heading("2. DETALLE DE PRECIO OFERTA BASE", level=1)
        h2.runs[0].font.color.rgb = blue_color

        # Create Styled 6-Column Table
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_titles = [
            "Ítem",
            "Código Partida",
            "Descripción de la Partida",
            "Cant.",
            f"Precio Unit. Neto {currency}",
            f"Subtotal Venta {currency}"
        ]

        col_widths = [Inches(0.6), Inches(1.4), Inches(2.4), Inches(0.6), Inches(1.25), Inches(1.25)]

        for idx, text in enumerate(hdr_titles):
            hdr_cells[idx].text = text
            set_cell_background(hdr_cells[idx], "1E293B")
            hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
            hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, line in enumerate(lines, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = str(line.get("item_code", ""))
            row_cells[2].text = str(line.get("name") or line.get("description", ""))
            row_cells[3].text = str(line.get("product_uom_qty", 1))
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            price_unit_val = float(line.get("price_unit", 0))
            price_subtotal_val = float(line.get("price_subtotal", 0))

            row_cells[4].text = format_currency(price_unit_val, currency)
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[5].text = format_currency(price_subtotal_val, currency)
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if idx % 2 == 0:
                for cell in row_cells:
                    set_cell_background(cell, "F8FAFC")

        # 3 Summary Rows at the end of summary table
        subtotal_row = table.add_row()
        subtotal_row.cells[4].text = "Subtotal Venta Neto"
        subtotal_row.cells[4].paragraphs[0].runs[0].font.bold = True
        subtotal_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        subtotal_row.cells[5].text = format_currency(amount_untaxed, currency)
        subtotal_row.cells[5].paragraphs[0].runs[0].font.bold = True
        subtotal_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_background(subtotal_row.cells[4], "F1F5F9")
        set_cell_background(subtotal_row.cells[5], "F1F5F9")

        iva_row = table.add_row()
        iva_row.cells[4].text = "IVA 19%"
        iva_row.cells[4].paragraphs[0].runs[0].font.bold = True
        iva_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        iva_row.cells[5].text = format_currency(amount_tax, currency)
        iva_row.cells[5].paragraphs[0].runs[0].font.bold = True
        iva_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_background(iva_row.cells[4], "F1F5F9")
        set_cell_background(iva_row.cells[5], "F1F5F9")

        total_row = table.add_row()
        total_row.cells[4].text = f"Total General {currency}"
        total_row.cells[4].paragraphs[0].runs[0].font.bold = True
        total_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_row.cells[5].text = format_currency(amount_total, currency)
        total_row.cells[5].paragraphs[0].runs[0].font.bold = True
        total_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_background(total_row.cells[4], "E2E8F0")
        set_cell_background(total_row.cells[5], "E2E8F0")

        # Explicitly set column widths on all cells
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = col_widths[idx]

        # Summary Financial Paragraph
        p_fin = doc.add_paragraph()
        p_fin.paragraph_format.space_before = Pt(14)
        p_fin.paragraph_format.space_after = Pt(14)
        p_fin.add_run("RESUMEN FINANCIERO DE LA OFERTA:\n").bold = True
        p_fin.add_run(f"• Monto Neto Total (Sin IVA): {format_currency(amount_untaxed, currency)}\n").bold = True
        p_fin.add_run(f"• Impuesto IVA (19%): {format_currency(amount_tax, currency)}\n")
        p_fin.add_run(f"• Monto Total Bruto (Con IVA): {format_currency(amount_total, currency)}\n").bold = True
        p_fin.add_run(f"• Margen Bruto Retenido Target: {margin_pct:.1f}%\n")

        # 3. EXCLUSIONES DE LA OFERTA
        h3 = doc.add_heading("3. EXCLUSIONES DE LA OFERTA", level=1)
        h3.runs[0].font.color.rgb = blue_color
        p_exc = doc.add_paragraph()
        exclusions_text = payload.get("exclusions") or (
            "Se excluye expresamente cualquier suministro o actividad distinta a las mencionadas en esta propuesta, en particular:\n"
            "• Obras civiles mayores, canalizaciones subterráneas de fuerza o tendido de cables de AT/MT.\n"
            "• Suministro de relojes satelitales GPS adicionales si el cliente optó por usar infraestructura existente en subestación.\n"
            "• Modificaciones a la lógica de protecciones de marcas no especificadas en la lista de materiales."
        )
        p_exc.add_run(exclusions_text)

        # 4. VALIDEZ DE LA OFERTA
        h4 = doc.add_heading("4. VALIDEZ DE LA OFERTA", level=1)
        h4.runs[0].font.color.rgb = blue_color
        p_val = doc.add_paragraph()
        validity_text = payload.get("validity") or "La presente oferta técnico-económica tiene una validez de 30 días corridos a contar de la fecha de su presentación."
        p_val.add_run(validity_text)

        # 5. CONDICIONES DE PAGO
        h5 = doc.add_heading("5. CONDICIONES DE PAGO", level=1)
        h5.runs[0].font.color.rgb = blue_color
        p_pago = doc.add_paragraph()
        payment_text = payload.get("payment_conditions") or payload.get("payment_terms") or (
            "El pago del precio ofertado se efectuará mediante Estados de Pago (EDP) contra emisión de factura a 30 días:\n"
            "• EDP N° 1 (30% del Total Neto): A la recepción de la Orden de Compra (OC) oficial del cliente.\n"
            "• EDP N° 2 (40% del Total Neto): A la entrega y aprobación de la Ingeniería de Detalle y Suministros en taller.\n"
            "• EDP N° 3 (30% del Total Neto): A la Puesta en Servicio (SAT Terreno) y entrega del Informe IPES registrado ante el CEN."
        )
        p_pago.add_run(payment_text)

        # 6. TÉRMINOS Y CONDICIONES (T&C)
        h6 = doc.add_heading("6. TÉRMINOS Y CONDICIONES (T&C)", level=1)
        h6.runs[0].font.color.rgb = blue_color
        p_tc = doc.add_paragraph()
        tc_text = payload.get("terms_and_conditions") or payload.get("terms") or (
            "• Garantía Técnica: 12 meses a contar de la recepción conforme de la puesta en servicio comercial.\n"
            "• Multas y Penalizaciones: Limitadas a un techo máximo acumulado del 5% del valor neto del contrato.\n"
            "• Fuerza Mayor: Suspensión de plazos en caso de eventos climáticos desfavorables en faena o indisponibilidad de la red eléctrica."
        )
        p_tc.add_run(tc_text)

        # Signature Block
        p_sig = doc.add_paragraph()
        p_sig.paragraph_format.space_before = Pt(30)
        p_sig.add_run(
            "Atentamente,\n\n\n"
            "_________________________________________\n"
            "CONECTA INGENIERÍA S.A.\n"
            "División de Operaciones & Propuestas Comercial OT\n"
            "comercial@conecta-ingenieria.cl | Santiago, Chile"
        ).bold = True

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream.getvalue()
